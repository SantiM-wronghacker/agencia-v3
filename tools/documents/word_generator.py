from datetime import datetime
from pathlib import Path

from tools.base import BaseTool, ToolResult, tool


@tool
class WordGeneratorTool(BaseTool):
    name = "word_generator"
    description = (
        "Genera documentos Word (.docx). "
        "Úsala para crear cartas, propuestas, "
        "contratos o documentos editables."
    )

    def run(self, action: str = "", **kwargs) -> ToolResult:
        if action == "generate":
            return self._generate(**kwargs)
        return self._error(f"Acción '{action}' no soportada")

    def _generate(
        self,
        title: str = "",
        content: str = "",
        output_path: str = None,
        company_name: str = "",
    ) -> ToolResult:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return self._error(
                "python-docx no instalado. Ejecuta: pip install python-docx"
            )

        if output_path is None:
            safe_title = title[:30].replace(" ", "_")
            output_path = f"data/docs/{safe_title}.docx"

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        if company_name:
            doc.add_heading(company_name, level=1)

        doc.add_heading(title, level=2)

        date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        date_para = doc.add_paragraph(f"Generado: {date_str}")
        date_para.runs[0].italic = True

        doc.add_paragraph("")  # spacer

        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        try:
            doc.save(output_path)
        except Exception as e:
            return self._error(f"Error al guardar documento Word: {e}")

        return self._success(
            f"Documento Word generado: {output_path}",
            raw_data={"path": output_path},
        )
